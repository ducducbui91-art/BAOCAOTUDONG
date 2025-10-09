# app.py
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import re
import os
import json
import zipfile
from typing import Dict
import io
import smtplib, ssl
from email.message import EmailMessage
import google.generativeai as genai

# --- CẤU HÌNH BẢO MẬT ---
# Cố gắng lấy "secrets" khi deploy trên Streamlit Cloud
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
    APP_EMAIL = st.secrets["APP_EMAIL"]
    APP_PASSWORD = st.secrets["APP_PASSWORD"]
# Nếu không được (chạy ở máy local), dùng các giá trị bên dưới
except Exception:
    st.warning("Không tìm thấy Streamlit Secrets. Đang sử dụng cấu hình local. Đừng quên thiết lập Secrets khi deploy!")
    # !!! QUAN TRỌNG: THAY THẾ CÁC GIÁ TRỊ DƯỚI ĐÂY BẰNG THÔNG TIN CỦA BẠN !!!
    GEMINI_API_KEY = "YOUR_GEMINI_API_KEY"
    APP_EMAIL = "your_email@example.com"
    APP_PASSWORD = "your_app_or_email_password"

# Cấu hình API key cho Gemini
try:
    genai.configure(api_key=GEMINI_API_KEY)
except Exception as e:
    st.error(f"Lỗi cấu hình Gemini API: {e}. Vui lòng kiểm tra lại API Key.")


#======================================================================
# PHẦN 1: ĐỊNH NGHĨA CÁC HÀM XỬ LÝ (HỘP CÔNG CỤ)
#======================================================================

def extract_vars_and_desc(docx_file_buffer) -> Dict[str, str]:
    """Trích xuất các biến và mô tả từ một file docx trong bộ nhớ."""
    xml_parts = []
    with zipfile.ZipFile(docx_file_buffer) as z:
        for name in z.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml_parts.append(z.read(name).decode("utf8"))
    all_xml = "\\n".join(xml_parts)
    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", all_xml, flags=re.DOTALL)
    full_text = "".join(texts)
    pattern = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}\s*\{#\s*(.*?)\s*#\}", flags=re.DOTALL)
    return dict(pattern.findall(full_text))

def call_gemini_model(transcript_content, placeholders):
    """Gửi yêu cầu đến Gemini và nhận về kết quả JSON."""
    model = genai.GenerativeModel("gemini-2.5-flash")
    Prompt_word ="""
# Vai trò
Bạn là một trợ lý AI chuyên nghiệp, có nhiệm vụ trích xuất thông tin quan trọng từ bản ghi cuộc họp để tạo ra nội dung cho biên bản cuộc họp, đảm bảo tính chính xác và trình bày chuyên nghiệp.

# Đầu vào
1.  **Bản ghi cuộc họp (transcript):** `{0}`
2.  **Danh sách các trường thông tin cần trích xuất (placeholders):** `{1}` (Đây là một đối tượng/dictionary nơi mỗi key là tên trường cần trích xuất và value là mô tả/yêu cầu định dạng cho trường đó).

# Nhiệm vụ
1.  **Phân tích kỹ lưỡng:** Đọc và hiểu toàn bộ nội dung bản ghi cuộc họp.
2.  **Xác định và Trích xuất:** Với **từng** trường thông tin (key) trong danh sách `placeholders`:
    *   Tìm (các) phần nội dung tương ứng trong bản ghi.
    *   Trích xuất thông tin một cách **chi tiết, đầy đủ ý, và chính xác tuyệt đối** về mặt ngữ nghĩa so với bản ghi gốc.
    *   **Trường hợp không có thông tin:** Nếu không tìm thấy thông tin rõ ràng cho một trường cụ thể trong bản ghi, hãy ghi nhận là "Chưa có thông tin".
3.  **Định dạng và Diễn đạt:**
    *   **Luôn trả về bằng tiếng Việt.**
    *   Sử dụng **văn phong trang trọng, lịch sự, chuyên nghiệp**, phù hợp với tiêu chuẩn của một biên bản cuộc họp chính thức.
    *   Diễn đạt thành **câu văn hoàn chỉnh, mạch lạc, đúng ngữ pháp và chính tả tiếng Việt**. Tổng hợp các ý rời rạc hoặc văn nói thành cấu trúc văn viết chuẩn mực.
    *   Đảm bảo mỗi thông tin trích xuất đều **rõ ràng, súc tích và có ý nghĩa**.
    *   **Quan trọng:** Áp dụng **đúng định dạng trình bày** (ví dụ: bullet cấp 1, bullet cấp 2, bảng Markdown, đoạn văn...) **theo yêu cầu được chỉ định trong phần mô tả (value) của placeholder tương ứng**.
4.  **Tạo đối tượng JSON:** Tập hợp tất cả thông tin đã trích xuất và định dạng vào một đối tượng JSON duy nhất, tuân thủ nghiêm ngặt các quy tắc xuất kết quả.

# Quy tắc xuất kết quả (Quan trọng - Tuân thủ nghiêm ngặt)
1.  **Khóa (keys) của JSON:**
    *   Phải **trùng khớp 100%** với từng phần tử (key) trong danh sách `placeholders`.
    *   Giữ nguyên mọi ký tự: dấu, dấu câu, khoảng trắng, chữ hoa/thường.
    *   **Tuyệt đối không:** chuyển sang không dấu, snake_case, camelCase, viết tắt, hoặc thay đổi tên khóa.
2.  **Cấu trúc JSON:**
    *   Chỉ xuất các cặp key-value tương ứng với `placeholders`.
    *   **Không** thêm khóa mới, **không** bớt khóa, **không** lồng ghép cấu trúc khác.
3.  **Giá trị (values) của JSON:**
    *   **Tuân thủ Yêu cầu Định dạng từ Placeholder:** **Đây là điểm cực kỳ quan trọng.** Đối với **mỗi** trường thông tin (key) trong JSON, bạn phải **đọc kỹ yêu cầu định dạng được nêu trong phần mô tả (value) của placeholder tương ứng** trong danh sách `placeholders`. **Áp dụng chính xác** định dạng đó cho chuỗi giá trị (value) của trường đó.
        *   Ví dụ: Nếu placeholder có yêu cấu trình bày theo bullet cấp 2 thì giá trị value trong Json phải bắt đầu mỗi dòng bằng '+'; hoặc nếu placeholder yêu cầu trình bày là dạng bảng thì giá trị key trong Json phải bắt buộc là dạng bảng markdown.
    *   **Nội dung:** Phải là kết quả đã được xử lý theo **Mục 3 (Định dạng và Diễn đạt)** ở phần Nhiệm vụ, đồng thời được **trình bày một cách rõ ràng, có cấu trúc chặt chẽ, và chuyên nghiệp** theo đúng yêu cầu định dạng từ placeholder.
    *   **Kiểu dữ liệu:** Tất cả giá trị (values) trong JSON phải là kiểu **chuỗi (string)**. **Tuyệt đối không sử dụng kiểu mảng (array) hoặc các kiểu dữ liệu khác.**
    *   **Xử lý trường hợp không có thông tin:** Nếu không tìm thấy thông tin cho một trường cụ thể trong bản ghi, giá trị tương ứng trong JSON phải là chuỗi: `Chưa có thông tin`.
    *   **Hướng dẫn Định dạng Bullet (KHI được yêu cầu trong Placeholder):** Mục tiêu là tạo ra văn bản có cấu trúc, dễ đọc và chuyên nghiệp. **Toàn bộ cấu trúc này phải được thể hiện bên trong chuỗi giá trị.**
        *   **Bullet cấp 1 (Thường dùng cho mục chính):** Bắt đầu dòng bằng dấu gạch ngang theo sau là một khoảng trắng (`- `) cho mỗi ý chính.
        *   **Bullet cấp 2 (Thường dùng cho ý phụ, chi tiết):** Bắt đầu dòng bằng dấu cộng theo sau là một khoảng trắng (`+ `) cho mỗi ý phụ. Nên thụt lề đầu dòng cho các mục cấp 2 (ví dụ: thêm 2 hoặc 4 dấu cách trước dấu `+ `) để phân biệt rõ ràng với cấp 1.
        *   **Trình bày dòng:** Mỗi mục bullet (cả `- ` và `+ `) phải nằm trên một dòng riêng biệt trong chuỗi kết quả. AI cần đảm bảo việc xuống dòng phù hợp giữa các mục bullet để tạo cấu trúc danh sách rõ ràng khi chuỗi được hiển thị.
        *   *Ví dụ cấu trúc bullet bên trong chuỗi giá trị (nếu placeholder yêu cầu `-` cho cấp 1 và `+` cho cấp 2):*
            ```
            - [Nội dung mục cấp 1 thứ nhất]
            - [Nội dung mục cấp 1 thứ hai]
              + [Nội dung mục cấp 2.1 thuộc mục 1.2]
              + [Nội dung mục cấp 2.2 thuộc mục 1.2]
            - [Nội dung mục cấp 1 thứ ba]
            ```
        *   **Đặc biệt với Công việc cần làm (Action Items) (NẾU placeholder yêu cầu cấu trúc này):** Cấu trúc rõ ràng thông tin cho từng mục, ví dụ sử dụng bullet cấp 1 (`- `) cho mỗi công việc và bullet cấp 2 (`+ `) thụt lề cho các chi tiết:
            ```
            - [Nội dung công việc cụ thể 1]
              + Người phụ trách: [Tên người/Bộ phận]
              + Hạn chót: [Ngày/Thời hạn cụ thể]
            - [Nội dung công việc cụ thể 2]
              + Người phụ trách: [Tên người/Bộ phận]
              + Hạn chót: [Ngày/Thời hạn cụ thể]
            ```
        *   **Tính nhất quán:** Áp dụng định dạng (bullet, bảng, đoạn văn...) một cách nhất quán theo đúng yêu cầu của từng placeholder.
4.  **Định dạng đầu ra:**
    *   **Không** bao gồm bất kỳ chú thích, giải thích, lời dẫn nào bên ngoài đối tượng JSON (ví dụ: không có `Đây là kết quả:` hay ```json ... ```).
    *   Toàn bộ kết quả trả về phải là **một chuỗi JSON hợp lệ và duy nhất**.
    """
    prompt = Prompt_word.format(transcript_content, placeholders)
    try:
        response = model.generate_content(
            contents=prompt,
            generation_config={"response_mime_type": "application/json"}
        )
        if response and hasattr(response, "text"):
            raw = response.text.strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1].strip("json\n")
            return json.loads(raw)
        else:
            st.error("Phản hồi từ Gemini API bị thiếu hoặc không hợp lệ.")
            return None
    except Exception as e:
        st.error(f"Lỗi khi gọi Gemini API: {e}")
        return None

# --- Khối code dài để xử lý file Word ---
COMMENT_RE     = re.compile(r"\{#.*?#\}")
COMMENT_ALL_RE = re.compile(r"\{#.*?#\}", re.DOTALL)
BOLD_RE        = re.compile(r"\*\*(.*?)\*\*")
TOKEN_RE       = re.compile(r"\{\{([^{}]+)\}\}")

def _is_md_table(text: str) -> bool:
    lines = [l.strip() for l in (text or "").strip().splitlines() if l.strip()]
    return len(lines) >= 2 and "|" in lines[0] and set(lines[1].replace(" ", "").replace(":", "")) <= set("-|")

def _parse_md_table(text: str):
    lines  = [l.strip() for l in (text or "").strip().splitlines() if l.strip()]
    header = [c.strip() for c in lines[0].split("|") if c.strip()]
    rows   = []
    for ln in lines[2:]:
        cols = [c.strip() for c in ln.split("|") if c.strip()]
        if cols:
            rows.append(cols)
    return header, rows

def _insert_table_after(paragraph, header, rows):
    # ... (Code chèn bảng của bạn) ...
    pass

def replace_in_paragraph(paragraph, data):
    # ... (Toàn bộ logic hàm replace_in_paragraph của bạn) ...
    pass

def fill_template_to_buffer(template_file_buffer, data_input):
    """Điền dữ liệu vào template và trả về file Word trong bộ nhớ (buffer)."""
    try:
        doc = Document(template_file_buffer)

        # Body
        for para in doc.paragraphs:
            replace_in_paragraph(para, data_input)
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, data_input)

        # Lưu file vào bộ nhớ
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer
    except Exception as e:
        st.error(f"Đã xảy ra lỗi khi tạo file Word: {e}")
        return None

def send_email_with_attachment(recipient_email, attachment_buffer, filename="BBCH.docx"):
    """Gửi email với file đính kèm từ buffer."""
    SMTP_SERVER = "smtp.office365.com"
    SMTP_PORT = 587
    
    msg = EmailMessage()
    msg["Subject"] = "Biên bản cuộc họp đã được tạo tự động"
    msg["From"] = APP_EMAIL
    msg["To"] = recipient_email
    msg.set_content(f"Chào bạn,\n\nBiên bản cuộc họp đã được tạo thành công.\nVui lòng xem trong file đính kèm.\n\nTrân trọng,\nCông cụ tạo biên bản tự động.")

    msg.add_attachment(
        attachment_buffer.getvalue(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
    
    try:
        ctx = ssl.create_default_context()
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as s:
            s.starttls(context=ctx)
            s.login(APP_EMAIL, APP_PASSWORD)
            s.send_message(msg)
        return True
    except Exception as e:
        st.error(f"Lỗi khi gửi email: {e}. Vui lòng kiểm tra lại cấu hình email và mật khẩu ứng dụng.")
        return False

#======================================================================
# PHẦN 2: GIAO DIỆN VÀ LUỒNG CHẠY CỦA WEBAPP (CẬP NHẬT)
#======================================================================

st.set_page_config(layout="wide", page_title="Công cụ tạo Biên bản cuộc họp")
st.title("🛠️ Công cụ tạo biên bản cuộc họp tự động")

with st.sidebar:
    st.info("📝 **Hướng dẫn sử dụng**")
    st.markdown("""
    1.  **Tải file transcript:** Tải lên file `.docx` chứa nội dung cuộc họp.
    2.  **Chọn Template:**
        * Sử dụng mẫu có sẵn bằng cách chọn "Template VPI".
        * Sử dụng mẫu riêng bằng cách chọn "Template tùy chỉnh" và tải file của bạn lên.
    3.  **Điền thông tin:** Nhập các thông tin cơ bản của cuộc họp.
    4.  **Nhập email:** Điền địa chỉ email bạn muốn nhận kết quả.
    5.  **Chạy:** Nhấn nút 'Tạo biên bản' và chờ trong giây lát.
    """)
    st.markdown("---")
    st.success("Ứng dụng được phát triển bởi VPI.")

st.header("📌 Nhập thông tin đầu vào")

transcript_file = st.file_uploader("1. Tải lên file transcript (.docx)", type=["docx"])

st.subheader("2. Lựa chọn Template")
template_option = st.selectbox(
    "Bạn muốn sử dụng loại template nào?",
    ("Template VPI", "Template tùy chỉnh"),
    help="Chọn 'Template VPI' để dùng mẫu có sẵn hoặc 'Template tùy chỉnh' để tải lên file của riêng bạn."
)

template_file = None
# Chỉ hiện ô upload khi người dùng chọn "Template tùy chỉnh"
if template_option == "Template tùy chỉnh":
    template_file = st.file_uploader("Tải lên file template .docx của bạn", type=["docx"])

st.subheader("3. Thông tin cơ bản")
col1, col2 = st.columns(2)
with col1:
    meeting_name = st.text_input("Tên cuộc họp")
    meeting_time = st.text_input("Thời gian cuộc họp (VD: 10/9/2025)")
    meeting_location = st.text_input("Địa điểm cuộc họp")
with col2:
    meeting_chair = st.text_input("Tên chủ trì")
    meeting_secretary = st.text_input("Tên thư ký")

recipient_email = st.text_input("4. Email nhận kết quả")

# Khi người dùng nhấn nút này, toàn bộ code xử lý MỚI BẮT ĐẦU CHẠY
if st.button("🚀 Tạo biên bản", type="primary"):
    
    # --- Bắt đầu luồng xử lý ---
    
    # Bước 1: Kiểm tra các đầu vào cơ bản
    if not all([transcript_file, recipient_email, meeting_name]):
        st.warning("Vui lòng tải lên file transcript và điền đầy đủ Tên cuộc họp, Email nhận kết quả.")
    else:
        # Bước 2: Xác định file template sẽ sử dụng
        template_to_use = None
        if template_option == "Template VPI":
            # Tên file template mặc định mà bạn đã đẩy lên GitHub
            template_to_use = "2025.VPI_BB hop 2025 1.docx" 
        elif template_file is not None:
            # File do người dùng tải lên
            template_to_use = template_file
        else:
            st.warning("Bạn đã chọn 'Template tùy chỉnh' nhưng chưa tải file lên.")

        # Bước 3: Nếu đã có đủ thông tin, bắt đầu xử lý
        if template_to_use:
            with st.spinner("⏳ Hệ thống đang xử lý, vui lòng chờ..."):
                try:
                    st.info("1/4 - Đang đọc và phân tích file...")
                    doc = Document(transcript_file)
                    transcript_content = "\\n".join([para.text for para in doc.paragraphs])
                    placeholders = extract_vars_and_desc(template_to_use)

                    st.info("2/4 - Đang gửi yêu cầu đến AI để tóm tắt...")
                    llm_result = call_gemini_model(transcript_content, placeholders)

                    if llm_result:
                        manual_inputs = {
                            'TenCuocHop': meeting_name, 'ThoiGianCuocHop': meeting_time,
                            'DiaDiemCuocHop': meeting_location, 'TenChuTri': meeting_chair,
                            'TenThuKy': meeting_secretary
                        }
                        llm_result.update(manual_inputs)

                        st.info("3/4 - Đang tạo file biên bản Word...")
                        docx_buffer = fill_template_to_buffer(template_to_use, llm_result)
                        
                        if docx_buffer:
                            st.info("4/4 - Đang gửi kết quả vào email của bạn...")
                            email_sent = send_email_with_attachment(recipient_email, docx_buffer)
                            
                            if email_sent:
                                st.success("✅ Hoàn thành! Biên bản sẽ được gửi tới email của bạn.")
                                st.balloons()
                            # (Thông báo lỗi gửi mail đã có trong hàm send_email_with_attachment)
                        else:
                             st.error("Không thể tạo file Word. Vui lòng kiểm tra lại file template.")
                    else:
                        st.error("Không thể lấy kết quả từ AI. Vui lòng thử lại.")
                except Exception as e:
                    st.error(f"Đã xảy ra một lỗi không mong muốn: {e}")
